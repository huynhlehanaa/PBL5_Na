from pathlib import Path
from preprocess import run_preprocess
from layout import run_layout
from ocr import run_ocr
from docx import Document

def main():
    # Cấu hình đường dẫn trỏ thẳng ra thư mục gốc (PBL5_Na)
    base = Path(__file__).resolve().parent.parent
    
    input_dir = base / "data" / "input"
    pre_layout_dir = base / "data" / "preprocessed" / "layout"
    pre_ocr_dir = base / "data" / "preprocessed" / "ocr"
    layout_dir = base / "data" / "layout"
    output_dir = base / "data" / "output"

    print("--- BƯỚC 1A: TIỀN XỬ LÝ CHO LAYOUT (IMUTILS) ---")
    run_preprocess(input_dir, pre_layout_dir, mode="layout")
    print("--- BƯỚC 1B: TIỀN XỬ LÝ CHO OCR (IMUTILS) ---")
    run_preprocess(input_dir, pre_ocr_dir, mode="ocr")

    print("\n--- BƯỚC 2: PHÂN TÍCH BỐ CỤC (DOCLAYOUT-YOLO) ---")
    layout_results = run_layout(pre_layout_dir, layout_dir)

    print("\n--- BƯỚC 3: NHẬN DIỆN CHỮ (VIETOCR) ---")
    ocr_results = run_ocr(pre_ocr_dir, layout_results, output_dir)

    print("\n--- BƯỚC 4: XUẤT FILE WORD ---")
    doc = Document()
    for res in ocr_results:
        doc.add_heading(f"Kết quả: {res['image']}", level=1)
        for line in res['content']:
            doc.add_paragraph(line)
    
    # Tạo thư mục output nếu chưa có và lưu file Word
    output_dir.mkdir(parents=True, exist_ok=True)
    doc_path = output_dir / "Ket_qua_PBL5.docx"
    doc.save(str(doc_path))
    print(f"✅ Hoàn thành! File Word tại: {doc_path}")

if __name__ == "__main__":
    main()